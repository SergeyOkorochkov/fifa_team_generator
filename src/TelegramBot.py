# работа с файловой системой (например, для загрузки переменных окружения).
import os

import random

# для работы с датой и временем (создание уникального имени файла).
from datetime import datetime

# from telegram.ext import (ApplicationBuilder, CommandHandler, MessageHandler, filters, ContextTypes, ConversationHandler)
from telegram import Update, InputFile
from telegram.ext import ApplicationBuilder
from telegram.ext import CommandHandler
from telegram.ext import MessageHandler
from telegram.ext import filters
from telegram.ext import ContextTypes
from telegram.ext import ConversationHandler


from docx import Document


class TelegramBot(object):
    __team_names = None
    __positions = None
    __positions_count = None
    __token = None

    def __init__(self, team_names, positions, positions_count, token):
        self.__team_names = team_names
        self.__positions = positions
        self.__positions_count = positions_count
        self.__token = token

    async def __start(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        await update.message.reply_text("Привет! Введи название первой команды:")
        return self.__team_names

    async def __cancel(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        await update.message.reply_text("Диалог отменён. Начни заново с /start.")
        return ConversationHandler.END  # Завершаем диалог

    async def team_names(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        if "team_1" not in context.user_data:  # проверка, это словарь???????????????????????? не понял
            context.user_data["team_1"] = update.message.text.strip()
            await update.message.reply_text("Теперь введи название второй команды:")
            return self.__team_names

        context.user_data["team_2"] = update.message.text.strip()
        context.user_data[
            "positions_count"] = {}  # создаёт пустой словарь для сохранения количества игроков на каждой позиции
        context.user_data["current_position"] = list(self.__positions.keys())[0]  # извлекает ключ 0 (Вратари)

        # Важное исправление! Здесь считаем максимальное количество для текущей позиции (вратари)
        position = context.user_data["current_position"]  # возвращает вратарей
        max_count = len(self.__positions[position]) // 2  # Максимальное количество игроков на позицию
        await update.message.reply_text(f"Сколько {position.lower()} в каждой команде? (максимум {max_count})")
        return self.__positions_count

    # Вверху и внизу ОДИНАКОВЫЙ код чтоли??????????????????????????????

    async def positions_count(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        position = context.user_data["current_position"]
        max_count = len(self.__positions[position]) // 2  # Максимальное количество игроков на текущую позицию
        try:
            count = int(update.message.text)
            if 0 <= count <= max_count:
                context.user_data["positions_count"][
                    position] = count  # сохраняем колличество игроков, если все без ошибок
                keys = list(self.__positions.keys())  # создаем список всех позиций? НО уже создавали ведь
                current_index = keys.index(position)
                if current_index + 1 < len(keys):
                    context.user_data["current_position"] = keys[current_index + 1]
                    next_position = context.user_data["current_position"]
                    next_max_count = len(self.__positions[next_position]) // 2
                    await update.message.reply_text(
                        f"Сколько {next_position.lower()} в каждой команде? (максимум {next_max_count})")
                    return self.__positions_count

                return await self.__process_teams(update, context)
            else:
                await update.message.reply_text(
                    f"Ошибка: нельзя выбрать больше {max_count} игроков на позицию {position}!")
                return self.__positions_count  # <-- Добавляем возврат состояния после ошибки
        except ValueError:
            await update.message.reply_text("Введите целое число!")
        return self.__positions_count

    async def __process_teams(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        team_1, team_2 = {}, {}
        for position, players in self.__positions.items():
            random.shuffle(players)
            count = context.user_data["positions_count"][position]  # извлекается количество игроков
            selected_players = players[:count * 2]  # ????????????????????????????????????????
            team_1[position], team_2[position] = selected_players[::2], selected_players[1::2]  # все четные и нечетные

        doc = Document()
        doc.add_heading('Состав команд', 0)

        def __add_players_table(doc, team_name, players):
            doc.add_heading(team_name, level=2)
            table = doc.add_table(rows=0, cols=2)  # таблица будет из 2 столбцов
            table.style = 'Table Grid'  # стиль с видимыми границами
            for i in range(0, len(players), 2):  # Проходим по списку игроков с шагом 2
                row = table.add_row().cells  # Добавляем новую строку в таблицу
                row[0].text = players[i]  # Первый игрок в первой ячейке
                if i + 1 < len(players):  # Если есть второй игрок, записываем во вторую ячейку
                    row[1].text = players[i + 1]

        for team_name, team in [(context.user_data["team_1"], team_1), (context.user_data["team_2"], team_2)]:
            doc.add_heading(team_name, level=1)
            for position, players in team.items():
                __add_players_table(doc, position, players)  # Добавляем таблицу с игроками по позициям????????????
            doc.add_page_break()  # Разрыв страницы

        file_path = f"teams_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx"  # уникальная дата
        doc.save(file_path)

        with open(file_path, "rb") as file:  # Открываем файл в бинарном ("rb") режиме.????????
            await update.message.reply_document(document=InputFile(file),
                                                filename=file_path)  # Отправляем его в чат с помощью reply_document
        os.remove(file_path)  # После отправки файл удаляется
        return ConversationHandler.END

    def run(self):
        app = (ApplicationBuilder()
               .token(self.__token)
               .build())  # создаем приложение бота???????????????

        conv_handler = ConversationHandler(
            entry_points=[CommandHandler('start', self.__start)],  # Начинаем с команды /start
            states={
                self.__team_names: [MessageHandler(filters.TEXT & ~filters.COMMAND, self.__team_names)],
                self.__positions_count: [MessageHandler(filters.TEXT & ~filters.COMMAND, self.__positions_count)],
            },
            fallbacks=[CommandHandler("cancel", self.__cancel)]
        )
        app.add_handler(conv_handler)
        app.run_polling()  # Запуск бота
        #  Если развернуть бота на сервере, лучше использовать вебхуки (webhook) вместо polling.??????

