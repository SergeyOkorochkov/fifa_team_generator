import random
from docx import Document
from datetime import datetime

# Игроки по позициям
positions = {
    "Вратари": ['Акинфеев', 'Ноер', 'Куртуа', 'Навас', 'Де Хеа', 'Тер Штеген', 'Облак', 'Буффон'],
    "Защитники": ['Рамос', 'Марсело', 'Карвахаль', 'Рафаэл', 'Варан', 'Тьяго Силва', 'Компани', 'Дани Алвес', 'Алаба',
                  'Маркос Алонсо', 'Рохо', 'Пике', 'Жорди Альба', 'Пепе', 'Лам', 'Валенсия', 'Смоллинг', 'Маскерано',
                  'Давид Луиз', 'Марио Фернандес', 'Филиппе Луиз', 'Беллерин', 'Боатенг', 'Ругани', 'Годин', 'Умтити',
                  'Сандро'],
    "Полузащитники": ['Иньеста', 'Модрич', 'Кросс', 'Конте', 'Вильям', 'Бускетс', 'Дембеле', 'Эриксен', 'Сон', 'Алли',
                      'Озил', 'Головин', 'Погба', 'Черышев', 'Ди Мария', 'Оскар', 'Алькантара', 'Де Брюйне', 'Ракитич',
                      'Иско', 'Марко Ройс', 'Хамес', 'Сильва', 'Фабрегас', 'Матюиди', 'Касорла', 'Гуардадо', 'Сане',
                      'Стерлинг'],
    "Нападающие": ['Азар', 'Неймар', 'Роналду', 'Дибала', 'Мбаппе', 'Кавани', 'Месси', 'Коутиньо', 'Дзюба', 'Дембеле',
                   'Суарес', 'Агуэро', 'Левандовский', 'Кейн', 'Гризман', 'Лукаку', 'Салах', 'Фирмино', 'Промес',
                   'Коста', 'Бэйл', 'Мане', 'Санчес', 'Икарди', 'Аубемейанг', 'Ляказетт', 'Джеко']
}

# Получаем названия команд
team_1_name = input("Введите название первой команды: ")
team_2_name = input("Введите название второй команды: ")

# Запрос количества игроков на каждую позицию
positions_count = {}
for position in positions.keys():
    while True:
        try:
            count = int(input(f"Сколько {position.lower()} вы хотите в каждой команде? "))
            if count >= 0:
                positions_count[position] = count
                break
            else:
                print("Число должно быть положительным!")
        except ValueError:
            print("Введите целое число!")

# Функция для распределения игроков
def divide_players(players, count):
    random.shuffle(players)
    selected_players = players[:count * 2]  # Оставляем только нужное количество игроков
    return selected_players[::2], selected_players[1::2]

# Генерация команд
team_1 = {}
team_2 = {}

for position, players in positions.items():
    team_1[position], team_2[position] = divide_players(players, positions_count[position])

# Создаём документ Word
doc = Document()
doc.add_heading('Состав команд', 0)

# Функция для добавления игроков в таблицу
def add_players_table(doc, team_name, players):
    doc.add_heading(team_name, level=2)
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    for i in range(0, len(players), 2):
        row = table.add_row().cells
        row[0].text = players[i]
        if i + 1 < len(players):
            row[1].text = players[i + 1]

# Добавление команд в документ
for team_name, team in [(team_1_name, team_1), (team_2_name, team_2)]:
    doc.add_heading(team_name, level=1)
    for position, players in team.items():
        add_players_table(doc, position, players)
    doc.add_page_break()

# Генерируем имя файла и сохраняем
file_name = f"Составы_команд_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx"

try:
    doc.save(file_name)
    print(f"Состав команд сохранен в файл {file_name}")
except Exception as e:
    print(f"Ошибка при сохранении файла: {e}")
