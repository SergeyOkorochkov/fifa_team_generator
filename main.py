import os #работа с файловой системой (например, для загрузки переменных окружения).
import random
from datetime import datetime #для работы с датой и временем (создание уникального имени файла).
from dotenv import load_dotenv  # Загружаем библиотеку для работы с .env
from telegram import Update, InputFile
from telegram.ext import (ApplicationBuilder, CommandHandler,
                          MessageHandler, filters, ContextTypes, ConversationHandler)
from docx import Document

# Загружаем переменные окружения
load_dotenv()

# Получаем токен
TOKEN = os.getenv("TOKEN")

# Проверяем, загружен ли токен, если нет выдаст ошибку специально
if not TOKEN:
    raise ValueError("Ошибка: Токен не найден! Проверь .env-файл.")

print("Токен успешно загружен!")


# Состояния для бота
TEAM_NAMES, POSITIONS_COUNT, PROCESSING = range(3) #Здесь создаются константы, которые будут использоваться
# в ConversationHandler (обработчике состояний).
# Игроки по позициям
positions = {
    "Вратари": ['Акинфеев', 'Ноер', 'Куртуа', 'Навас', 'Де Хеа',
                'Тер Штеген', 'Облак', 'Буффон'],
    "Защитники": ['Рамос', 'Марсело', 'Карвахаль', 'Рафаэл',
                  'Варан', 'Тьяго Силва', 'Компани', 'Дани Алвес',
                  'Алаба', 'Маркос Алонсо', 'Рохо', 'Пике',
                  'Жорди Альба', 'Пепе', 'Лам', 'Валенсия', 'Смоллинг',
                  'Маскерано', 'Давид Луиз', 'Марио Фернандес',
                  'Филиппе Луиз', 'Беллерин', 'Боатенг', 'Ругани',
                  'Годин', 'Умтити', 'Сандро'],
    "Полузащитники": ['Иньеста', 'Модрич', 'Кросс', 'Конте',
                      'Вильям', 'Бускетс', 'Дембеле', 'Эриксен', 'Сон', 'Алли',
                      'Озил', 'Головин', 'Погба', 'Черышев',
                      'Ди Мария', 'Оскар', 'Алькантара', 'Де Брюйне', 'Ракитич',
                      'Иско', 'Марко Ройс', 'Хамес', 'Сильва',
                      'Фабрегас', 'Матюиди', 'Касорла', 'Гуардадо', 'Сане',
                      'Стерлинг'],
    "Нападающие": ['Азар', 'Неймар', 'Роналду', 'Дибала', 'Мбаппе',
                   'Кавани', 'Месси', 'Коутиньо', 'Дзюба', 'Дембеле',
                   'Суарес', 'Агуэро', 'Левандовский', 'Кейн',
                   'Гризман', 'Лукаку', 'Салах', 'Фирмино', 'Промес',
                   'Коста', 'Бэйл', 'Мане', 'Санчес', 'Икарди',
                   'Аубемейанг', 'Ляказетт', 'Джеко']
}


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Привет! Введи название первой команды:")
    return TEAM_NAMES

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Диалог отменён. Начни заново с /start.")
    return ConversationHandler.END  # Завершаем диалог


async def team_names(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if "team_1" not in context.user_data: #проверка, это словарь???????????????????????? не понял
        context.user_data["team_1"] = update.message.text.strip()
        await update.message.reply_text("Теперь введи название второй команды:")
        return TEAM_NAMES

    context.user_data["team_2"] = update.message.text.strip()
    context.user_data["positions_count"] = {} #создаёт пустой словарь для сохранения количества игроков на каждой позиции
    context.user_data["current_position"] = list(positions.keys())[0] # извлекает ключ 0 (Вратари)

    # Важное исправление! Здесь считаем максимальное количество для текущей позиции (вратари)
    position = context.user_data["current_position"] #возвращает вратарей
    max_count = len(positions[position]) // 2  # Максимальное количество игроков на позицию
    await update.message.reply_text(f"Сколько {position.lower()} в каждой команде? (максимум {max_count})")
    return POSITIONS_COUNT

# Вверху и внизу ОДИНАКОВЫЙ код чтоли??????????????????????????????

async def positions_count(update: Update, context: ContextTypes.DEFAULT_TYPE):
    position = context.user_data["current_position"]
    max_count = len(positions[position]) // 2  # Максимальное количество игроков на текущую позицию
    try:
        count = int(update.message.text)
        if 0 <= count <= max_count:
            context.user_data["positions_count"][position] = count #сохраняем колличество игроков, если все без ошибок
            keys = list(positions.keys()) #создаем список всех позиций? НО уже создавали ведь
            current_index = keys.index(position)
            if current_index + 1 < len(keys):
                context.user_data["current_position"] = keys[current_index + 1]
                next_position = context.user_data["current_position"]
                next_max_count = len(positions[next_position]) // 2
                await update.message.reply_text(
                    f"Сколько {next_position.lower()} в каждой команде? (максимум {next_max_count})")
                return POSITIONS_COUNT

            return await process_teams(update, context)
        else:
            await update.message.reply_text(f"Ошибка: нельзя выбрать больше {max_count} игроков на позицию {position}!")
            return POSITIONS_COUNT  # <-- Добавляем возврат состояния после ошибки
    except ValueError:
        await update.message.reply_text("Введите целое число!")
    return POSITIONS_COUNT



async def process_teams(update: Update, context: ContextTypes.DEFAULT_TYPE):
    team_1, team_2 = {}, {}
    for position, players in positions.items():
        random.shuffle(players)
        count = context.user_data["positions_count"][position] #извлекается количество игроков
        selected_players = players[:count * 2] #????????????????????????????????????????
        team_1[position], team_2[position] = selected_players[::2], selected_players[1::2] #все четные и нечетные

    doc = Document()
    doc.add_heading('Состав команд', 0)

    def add_players_table(doc, team_name, players):
        doc.add_heading(team_name, level=2)
        table = doc.add_table(rows=0, cols=2) #таблица будет из 2 столбцов
        table.style = 'Table Grid' #стиль с видимыми границами
        for i in range(0, len(players), 2): # Проходим по списку игроков с шагом 2
            row = table.add_row().cells # Добавляем новую строку в таблицу
            row[0].text = players[i] # Первый игрок в первой ячейке
            if i + 1 < len(players): # Если есть второй игрок, записываем во вторую ячейку
                row[1].text = players[i + 1]

    for team_name, team in [(context.user_data["team_1"], team_1), (context.user_data["team_2"], team_2)]:
        doc.add_heading(team_name, level=1)
        for position, players in team.items():
            add_players_table(doc, position, players) # Добавляем таблицу с игроками по позициям????????????
        doc.add_page_break() # Разрыв страницы

    file_path = f"teams_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx" # уникальная дата
    doc.save(file_path)

    with open(file_path, "rb") as file: #Открываем файл в бинарном ("rb") режиме.????????
        await update.message.reply_document(document=InputFile(file), filename=file_path) #Отправляем его в чат с помощью reply_document
    os.remove(file_path) #После отправки файл удаляется
    return ConversationHandler.END


app = ApplicationBuilder().token(TOKEN).build() #создаем приложение бота???????????????
conv_handler = ConversationHandler(
    entry_points=[CommandHandler('start', start)], # Начинаем с команды /start
    states={
        TEAM_NAMES: [MessageHandler(filters.TEXT & ~filters.COMMAND, team_names)],
        POSITIONS_COUNT: [MessageHandler(filters.TEXT & ~filters.COMMAND, positions_count)],
    },
    fallbacks=[CommandHandler("cancel", cancel)]
)
app.add_handler(conv_handler)
app.run_polling() #Запуск бота
#  Если развернуть бота на сервере, лучше использовать вебхуки (webhook) вместо polling.??????